import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_word_guide():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles & Fonts
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(33, 37, 41)

    # Helper function for headings
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(139, 94, 60) # Milk tea brown
        p.paragraph_format.space_after = Pt(4)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        p.paragraph_format.space_after = Pt(24)

    def add_heading1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 26, 20)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)

    def add_heading2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(139, 94, 60)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    def add_callout(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("💡 MẸO VẼ TRÊN STARUML: " + text)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(25, 135, 84)

    # Document Header
    add_title("CẨM NANG HƯỚNG DẪN VẼ HANDS-ON 10 SƠ ĐỒ UML")
    add_subtitle("Dự Án: Quản Lý Bán Hàng Trà Sữa Cô Đào Quán POS | Công cụ: StarUML")

    # Introduction Box
    p_intro = doc.add_paragraph()
    run_intro = p_intro.add_run(
        "Tài liệu này được biên soạn chi tiết từng bước để giúp bạn TỰ VẼ MẤY TAY trên phần mềm StarUML. "
        "Với mỗi sơ đồ, tài liệu liệt kê chính xác: các Element (hình/đối tượng) cần kéo thả ra canvas, "
        "loại đường nối (Arrow type) và danh sách chi tiết điểm nối từ đâu đến đâu kèm nhãn tên."
    )
    run_intro.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ----------------------------------------------------
    # SƠ ĐỒ 1: USE CASE DIAGRAM
    # ----------------------------------------------------
    add_heading1("1. SƠ ĐỒ USE CASE TỔNG THỂ (Overall Use Case Diagram)")
    
    p = doc.add_paragraph()
    p.add_run("• Mục đích: ").bold = True
    p.add_run("Mô tả toàn bộ các chức năng của hệ thống POS chia theo 3 nhóm người dùng (Admin, Staff, Customer).")
    
    add_heading2("A. Tạo Diagram trong StarUML:")
    doc.add_paragraph("1. Vào Model -> Add Diagram -> Select 'Use Case Diagram'. Naming: 'UC_Overall_PosSystem'.")
    doc.add_paragraph("2. Vẽ một hình chữ nhật lớn đại diện cho Hệ thống (System Boundary), đặt tên: 'Hệ Thống Cô Đào Quán POS'.")

    add_heading2("B. Danh sách 3 Tác Nhân (Actors) cần kéo thả:")
    table_uc_actor = doc.add_table(rows=1, cols=3)
    table_uc_actor.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table_uc_actor.rows[0].cells
    hdr[0].text = "Tên Actor (StarUML Element)"
    hdr[1].text = "Vị trí đặt trên Canvas"
    hdr[2].text = "Vai trò"
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True

    actors_data = [
        ("Customer (Khách hàng)", "Bên trái Canvas (ngoài Boundary)", "Đặt món online, quét QR tại bàn, xem lịch sử & tích điểm"),
        ("Staff (Nhân viên / Thu ngân)", "Bên phải Canvas (ở trên)", "Bán hàng POS tại quầy, duyệt đơn online, quản lý bàn"),
        ("Admin (Quản lý)", "Bên phải Canvas (ở dưới)", "Quản trị toàn bộ danh mục, nguyên liệu, nhân viên, voucher, VietQR, báo cáo")
    ]
    for name, pos, role in actors_data:
        row = table_uc_actor.add_row().cells
        row[0].text = name
        row[1].text = pos
        row[2].text = role

    add_heading2("C. Danh sách 15 Use Case (Hình Elip) đặt bên trong System Boundary:")
    doc.add_paragraph("• Khách hàng: UC1 (Đăng ký/Đăng nhập), UC2 (Xem thực đơn online), UC3 (Thêm món vào giỏ), UC4 (Đặt đơn online), UC5 (Quét QR tại bàn), UC6 (Theo dõi đơn & Tích điểm).")
    doc.add_paragraph("• Nhân viên: UC7 (Bán hàng POS quầy 1-Click), UC8 (Thanh toán & In hóa đơn 80mm), UC9 (Quản lý bàn & Mã QR), UC10 (Duyệt / Hủy đơn online).")
    doc.add_paragraph("• Quản lý: UC11 (Quản lý Sản phẩm & Danh mục), UC12 (Quản lý Kho & Nguyên liệu), UC13 (Quản lý Nhân viên & Phân quyền), UC14 (Quản lý Voucher & VietQR), UC15 (Xem Báo cáo Doanh thu Excel).")

    add_heading2("D. Hướng dẫn nối đường (Connections & Relationships):")
    table_uc_conn = doc.add_table(rows=1, cols=4)
    table_uc_conn.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = table_uc_conn.rows[0].cells
    h[0].text = "Từ (Source)"
    h[1].text = "Đến (Target)"
    h[2].text = "Loại đường nối (StarUML Tool)"
    h[3].text = "Ghi chú / Nhãn"
    for cell in h:
        cell.paragraphs[0].runs[0].font.bold = True

    conns_uc = [
        ("Customer", "UC2, UC3, UC4, UC5, UC6", "Association (Đường thẳng nét liền)", "Nối Khách hàng với các UC chức năng của khách"),
        ("Staff", "UC1, UC7, UC8, UC9, UC10", "Association (Đường thẳng nét liền)", "Nối Nhân viên với các UC bán hàng & duyệt đơn"),
        ("Admin", "Staff", "Generalization (Đường mũi tên rỗng kế thừa)", "Admin kế thừa toàn bộ chức năng của Staff"),
        ("Admin", "UC11, UC12, UC13, UC14, UC15", "Association (Đường thẳng nét liền)", "Nối Admin với các UC quản trị kho & báo cáo"),
        ("UC7 (Bán hàng POS)", "UC8 (In Hóa Đơn)", "Include (Đường nét đứt <<include>>)", "Bán hàng xong bắt buộc bao gồm in hóa đơn"),
        ("UC4 (Đặt đơn online)", "UC6 (Tích điểm)", "Include (Đường nét đứt <<include>>)", "Hoàn tất đơn tự động tích điểm")
    ]
    for s, t, r, n in conns_uc:
        row = table_uc_conn.add_row().cells
        row[0].text = s
        row[1].text = t
        row[2].text = r
        row[3].text = n

    add_callout("Dùng công cụ 'Association' nối Actor tới Use Case. Dùng 'Generalization' kéo từ Admin trỏ mũi tên vào Staff!")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ----------------------------------------------------
    # SƠ ĐỒ 2: ERD DIAGRAM
    # ----------------------------------------------------
    add_heading1("2. SƠ ĐỒ THỰC THỂ MỐI QUAN HỆ CSDL (Full ERD Diagram)")
    
    p = doc.add_paragraph()
    p.add_run("• Mục đích: ").bold = True
    p.add_run("Mô tả cấu trúc CSDL gồm 17 thực thể (Entities), các trường thuộc tính (PK, FK) và quan hệ 1-n, 1-1.")
    
    add_heading2("A. Tạo Diagram trong StarUML:")
    doc.add_paragraph("Vào Model -> Add Diagram -> Select 'ERD Diagram' (hoặc Class Diagram dùng kiểu ERD). Naming: 'ERD_MilkTeaPos'.")

    add_heading2("B. Danh sách 17 Thực thể (Entities) cần tạo:")
    doc.add_paragraph("1. ROLE (role_id PK, role_name, description)")
    doc.add_paragraph("2. USERS (user_id PK, username, password, full_name, role_id FK)")
    doc.add_paragraph("3. CATEGORY (category_id PK, category_name, description)")
    doc.add_paragraph("4. PRODUCT (product_id PK, product_name, price, image, category_id FK)")
    doc.add_paragraph("5. INGREDIENT (ingredient_id PK, ingredient_name, quantity, unit)")
    doc.add_paragraph("6. PRODUCT_INGREDIENT (product_ingredient_id PK, product_id FK, ingredient_id FK, quantity_used)")
    doc.add_paragraph("7. INVENTORY_TRANSACTION (transaction_id PK, ingredient_id FK, quantity, transaction_type, transaction_date)")
    doc.add_paragraph("8. CUSTOMER (customer_id PK, full_name, phone, email, point)")
    doc.add_paragraph("9. TABLE_CAFE (table_id PK, table_number, status, qr_code)")
    doc.add_paragraph("10. VOUCHER (voucher_id PK, voucher_code, discount_percent, start_date, end_date)")
    doc.add_paragraph("11. ORDERS (order_id PK, order_date, total_amount, status, customer_id FK, table_id FK, voucher_id FK, user_id FK)")
    doc.add_paragraph("12. ORDER_DETAIL (order_detail_id PK, order_id FK, product_id FK, quantity, price, subtotal)")
    doc.add_paragraph("13. INVOICE (invoice_id PK, invoice_date, total_amount, order_id FK)")
    doc.add_paragraph("14. PAYMENT (payment_id PK, payment_method, payment_status, invoice_id FK)")
    doc.add_paragraph("15. CUSTOMER_ORDER (order_id PK, customer_id FK, order_date, total_amount, status)")
    doc.add_paragraph("16. CUSTOMER_ORDER_DETAIL (detail_id PK, order_id FK, product_id FK, quantity, price, subtotal)")
    doc.add_paragraph("17. BANK_SETTING (id PK, bank_id, account_no, account_name)")

    add_heading2("C. Hướng dẫn nối đường quan hệ (ERD Relationship Lines):")
    table_erd = doc.add_table(rows=1, cols=4)
    table_erd.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = table_erd.rows[0].cells
    h[0].text = "Thực thể A"
    h[1].text = "Thực thể B"
    h[2].text = "Tỷ lệ quan hệ (Cardinality)"
    h[3].text = "Loại đường nối trong StarUML"
    for cell in h:
        cell.paragraphs[0].runs[0].font.bold = True

    erd_conns = [
        ("ROLE", "USERS", "1 - N (1 Role có nhiều Users)", "One-to-Many Relationship"),
        ("CATEGORY", "PRODUCT", "1 - N (1 Danh mục chứa nhiều Sản phẩm)", "One-to-Many Relationship"),
        ("PRODUCT", "PRODUCT_INGREDIENT", "1 - N (1 Món có nhiều công thức nguyên liệu)", "One-to-Many Relationship"),
        ("INGREDIENT", "PRODUCT_INGREDIENT", "1 - N (1 Nguyên liệu dùng cho nhiều món)", "One-to-Many Relationship"),
        ("INGREDIENT", "INVENTORY_TRANSACTION", "1 - N (1 Nguyên liệu có nhiều nhật ký kho)", "One-to-Many Relationship"),
        ("CUSTOMER", "ORDERS", "1 - N (1 Khách có nhiều đơn hàng quầy)", "One-to-Many Relationship"),
        ("TABLE_CAFE", "ORDERS", "1 - N (1 Bàn gán cho nhiều đơn theo thời gian)", "One-to-Many Relationship"),
        ("VOUCHER", "ORDERS", "1 - N (1 Voucher áp dụng cho nhiều đơn)", "One-to-Many Relationship"),
        ("USERS", "ORDERS", "1 - N (1 Thu ngân lập nhiều đơn hàng)", "One-to-Many Relationship"),
        ("ORDERS", "ORDER_DETAIL", "1 - N (1 Đơn chứa nhiều dòng chi tiết món)", "One-to-Many Relationship"),
        ("PRODUCT", "ORDER_DETAIL", "1 - N (1 Sản phẩm bán trong nhiều dòng đơn)", "One-to-Many Relationship"),
        ("ORDERS", "INVOICE", "1 - 1 (1 Đơn xuất đúng 1 Hóa đơn)", "One-to-One Relationship"),
        ("INVOICE", "PAYMENT", "1 - 1 (1 Hóa đơn có 1 giao dịch thanh toán)", "One-to-One Relationship"),
        ("CUSTOMER", "CUSTOMER_ORDER", "1 - N (1 Khách đặt nhiều đơn online)", "One-to-Many Relationship"),
        ("CUSTOMER_ORDER", "CUSTOMER_ORDER_DETAIL", "1 - N (1 Đơn online chứa nhiều món)", "One-to-Many Relationship")
    ]
    for a, b, c, r in erd_conns:
        row = table_erd.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
        row[3].text = r

    add_callout("Dùng công cụ 'One-to-Many Relationship' kéo từ Bảng chứa PK (1) sang Bảng chứa FK (N)!")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ----------------------------------------------------
    # SƠ ĐỒ 3: CLASS DIAGRAM
    # ----------------------------------------------------
    add_heading1("3. SƠ ĐỒ LỚP KIẾN TRÚC 3 TẦNG (Class Diagram - 3-Tier Architecture)")
    
    p = doc.add_paragraph()
    p.add_run("• Mục đích: ").bold = True
    p.add_run("Mô tả kiến trúc 3 tầng MVC Spring Boot: Controller Layer, Service Layer, Repository Layer và Entity Layer.")
    
    add_heading2("A. Tạo Diagram trong StarUML:")
    doc.add_paragraph("Vào Model -> Add Diagram -> Select 'Class Diagram'. Naming: 'Class_Architecture_3Tier'.")

    add_heading2("B. Các Class chính cần tạo (Hình chữ nhật Class):")
    doc.add_paragraph("1. OrdersController (+addOrder(), +posCheckout())")
    doc.add_paragraph("2. OrdersService <<interface>> (+getAllOrders(), +saveOrder(), +getOrderById())")
    doc.add_paragraph("3. OrdersServiceImpl (+getAllOrders(), +saveOrder())")
    doc.add_paragraph("4. OrdersRepository <<interface>> (+findByStatus())")
    doc.add_paragraph("5. Orders Entity (-orderId, -orderDate, -totalAmount, -status)")
    doc.add_paragraph("6. InvoiceController (+viewInvoice())")
    doc.add_paragraph("7. InvoiceService <<interface>> (+saveInvoice(), +getInvoiceById())")
    doc.add_paragraph("8. InvoiceRepository <<interface>>")
    doc.add_paragraph("9. Invoice Entity (-invoiceId, -invoiceDate, -totalAmount)")

    add_heading2("C. Hướng dẫn nối đường liên kết giữa các Class:")
    table_class = doc.add_table(rows=1, cols=4)
    table_class.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = table_class.rows[0].cells
    h[0].text = "Class A"
    h[1].text = "Class B"
    h[2].text = "Mối quan hệ"
    h[3].text = "Loại đường trong StarUML Tool"
    for cell in h:
        cell.paragraphs[0].runs[0].font.bold = True

    class_conns = [
        ("OrdersController", "OrdersService", "Sử dụng Service", "Directed Association (Mũi tên nét liền trỏ vào Service)"),
        ("OrdersServiceImpl", "OrdersService", "Thực thi Interface", "Interface Realization (Mũi tên nét đứt trỏ vào Interface)"),
        ("OrdersServiceImpl", "OrdersRepository", "Gọi Repository", "Directed Association (Mũi tên nét liền trỏ vào Repo)"),
        ("OrdersRepository", "Orders Entity", "Thao tác trên Entity", "Dependency (Mũi tên nét đứt trỏ vào Entity)"),
        ("InvoiceController", "InvoiceService", "Sử dụng Service", "Directed Association (Mũi tên nét liền)"),
        ("InvoiceController", "OrdersService", "Sử dụng Service", "Directed Association (Mũi tên nét liền)")
    ]
    for a, b, c, r in class_conns:
        row = table_class.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
        row[3].text = r

    add_callout("Gõ <<interface>> ở mục Stereotype của các Lớp Interface!")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ----------------------------------------------------
    # SƠ ĐỒ 4: SEQUENCE DIAGRAM - POS COUNTER CHECKOUT
    # ----------------------------------------------------
    add_heading1("4. SƠ ĐỒ TUẦN TỰ - BÁN HÀNG TẠI QUẦY POS (Sequence Diagram - POS Counter Checkout)")
    
    p = doc.add_paragraph()
    p.add_run("• Mục đích: ").bold = True
    p.add_run("Mô tả tương tác từng bước từ khi Thu ngân bấm chọn món đến khi Hệ thống tự động trừ kho nguyên liệu, tích điểm và in hóa đơn.")

    add_heading2("A. Tạo Diagram trong StarUML:")
    doc.add_paragraph("Vào Model -> Add Diagram -> Select 'Sequence Diagram'. Naming: 'Seq_POS_Counter_Checkout'.")

    add_heading2("B. Các Lifelines (Cột đứng) xếp từ trái sang phải:")
    doc.add_paragraph("1. Staff (Actor: 👔 Thu ngân)")
    doc.add_paragraph("2. UI: Màn Hình POS (/orders/add)")
    doc.add_paragraph("3. Controller: OrdersController")
    doc.add_paragraph("4. OrderSvc: OrdersService")
    doc.add_paragraph("5. RecipeRepo: ProductIngredientRepository")
    doc.add_paragraph("6. StockRepo: IngredientRepository")
    doc.add_paragraph("7. CustRepo: CustomerRepository")
    doc.add_paragraph("8. InvoiceSvc: InvoiceService")
    doc.add_paragraph("9. PrintUI: Màn Hình In Hóa Đơn Nhiệt 80mm (/invoice/{id})")

    add_heading2("C. Danh sách các Thông điệp gửi đi (Sequence Messages):")
    table_seq1 = doc.add_table(rows=1, cols=5)
    table_seq1.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = table_seq1.rows[0].cells
    h[0].text = "STT"
    h[1].text = "Gửi từ (From)"
    h[2].text = "Đến (To)"
    h[3].text = "Nội dung Thông điệp (Message Text)"
    h[4].text = "Loại mũi tên (StarUML Tool)"
    for cell in h:
        cell.paragraphs[0].runs[0].font.bold = True

    seq1_data = [
        ("1", "Staff", "UI", "1: Chọn món uống, số bàn & bấm 'Thanh Toán POS'", "Message (Mũi tên nét liền)"),
        ("2", "UI", "Controller", "2: POST /orders/pos-checkout (productIds, qty, customerId)", "Message (Mũi tên nét liền)"),
        ("3", "Controller", "OrderSvc", "3: saveOrder(order)", "Message (Mũi tên nét liền)"),
        ("4", "OrderSvc", "Controller", "4: Return Order object (#1)", "Reply Message (Mũi tên nét đứt)"),
        ("5", "Controller", "RecipeRepo", "5: findByProductProductId(pId)", "Message (Mũi tên nét liền)"),
        ("6", "RecipeRepo", "Controller", "6: Return danh sách công thức nguyên liệu", "Reply Message (Mũi tên nét đứt)"),
        ("7", "Controller", "StockRepo", "7: Trừ tồn kho (Ingredient.quantity -= neededQty)", "Message (Mũi tên nét liền)"),
        ("8", "Controller", "StockRepo", "8: Save InventoryTransaction EXPORT", "Message (Mũi tên nét liền)"),
        ("9", "Controller", "CustRepo", "9: Cộng điểm thưởng (Customer.point += total/10k)", "Message (Mũi tên nét liền)"),
        ("10", "Controller", "InvoiceSvc", "10: saveInvoice(invoice) & savePayment(payment)", "Message (Mũi tên nét liền)"),
        ("11", "InvoiceSvc", "Controller", "11: Return Invoice object (#1)", "Reply Message (Mũi tên nét đứt)"),
        ("12", "Controller", "UI", "12: Redirect 302 /invoice/1", "Reply Message (Mũi tên nét đứt)"),
        ("13", "UI", "PrintUI", "13: GET /invoice/1", "Message (Mũi tên nét liền)"),
        ("14", "PrintUI", "Staff", "14: Hiển thị phiếu in nhiệt 80mm & Mã VietQR tự động", "Reply Message (Mũi tên nét đứt)")
    ]
    for stt, fr, to, msg, arrow in seq1_data:
        row = table_seq1.add_row().cells
        row[0].text = stt
        row[1].text = fr
        row[2].text = to
        row[3].text = msg
        row[4].text = arrow

    add_callout("Dùng 'Message' cho yêu cầu gọi hàm (nét liền), dùng 'Reply Message' cho kết quả trả về (nét đứt)!")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ----------------------------------------------------
    # SƠ ĐỒ 5: SEQUENCE DIAGRAM - KHÁCH ĐẶT ONLINE
    # ----------------------------------------------------
    add_heading1("5. SƠ ĐỒ TUẦN TỰ - KHÁCH ĐẶT MÓN ONLINE & ADMIN DUYỆT (Sequence Diagram - Online Order)")
    
    add_heading2("A. Lifelines kéo từ trái sang phải:")
    doc.add_paragraph("1. Customer (Actor: 🥤 Khách hàng)")
    doc.add_paragraph("2. CustomerUI: Customer Menu (/customer/menu)")
    doc.add_paragraph("3. CartCtrl: CustomerCartController")
    doc.add_paragraph("4. Admin (Actor: 👑 Admin / Thu ngân)")
    doc.add_paragraph("5. AdminUI: Trang Đơn Online (/admin/customer-orders)")
    doc.add_paragraph("6. AdminCtrl: AdminCustomerOrderController")

    add_heading2("B. Danh sách các Thông điệp (Sequence Messages):")
    table_seq2 = doc.add_table(rows=1, cols=5)
    table_seq2.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = table_seq2.rows[0].cells
    h[0].text = "STT"
    h[1].text = "Gửi từ"
    h[2].text = "Đến"
    h[3].text = "Nội dung Thông điệp"
    h[4].text = "Loại mũi tên"
    for cell in h:
        cell.paragraphs[0].runs[0].font.bold = True

    seq2_data = [
        ("1", "Customer", "CustomerUI", "1: Bấm 'Thêm vào giỏ hàng'", "Message"),
        ("2", "CustomerUI", "CartCtrl", "2: GET /customer-cart/add/{customerId}/{productId}", "Message"),
        ("3", "CartCtrl", "CustomerUI", "3: Cập nhật giỏ hàng CART", "Reply Message"),
        ("4", "Customer", "CustomerUI", "4: Bấm 'Đặt hàng Online'", "Message"),
        ("5", "CustomerUI", "CartCtrl", "5: GET /customer-order/checkout", "Message"),
        ("6", "CartCtrl", "CustomerUI", "6: Đơn chuyển trạng thái PENDING", "Reply Message"),
        ("7", "Admin", "AdminUI", "7: Xem danh sách đơn online chờ duyệt", "Message"),
        ("8", "AdminUI", "AdminCtrl", "8: GET /admin/customer-orders/approve/{id}", "Message"),
        ("9", "AdminCtrl", "AdminCtrl", "9: Tự động tạo Orders, trừ kho & cộng điểm", "Self Message (Mũi tên tự trỏ mình)"),
        ("10", "AdminCtrl", "AdminUI", "10: Cập nhật status COMPLETED & báo thành công", "Reply Message")
    ]
    for stt, fr, to, msg, arrow in seq2_data:
        row = table_seq2.add_row().cells
        row[0].text = stt
        row[1].text = fr
        row[2].text = to
        row[3].text = msg
        row[4].text = arrow

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ----------------------------------------------------
    # SƠ ĐỒ 6: SEQUENCE DIAGRAM - VIETQR & IN HÓA ĐƠN
    # ----------------------------------------------------
    add_heading1("6. SƠ ĐỒ TUẦN TỰ - VIETQR & IN HÓA ĐƠN 80MM (Sequence Diagram - VietQR & Print)")
    
    add_heading2("A. Lifelines kéo từ trái sang phải:")
    doc.add_paragraph("1. User (Actor: 👤 Khách hàng / Thu ngân)")
    doc.add_paragraph("2. InvoiceCtrl: InvoiceController")
    doc.add_paragraph("3. BankSvc: BankSettingService")
    doc.add_paragraph("4. VietQR: VietQR API (img.vietqr.io)")
    doc.add_paragraph("5. ThermalPage: Trang In Hóa Đơn Nhiệt 80mm")

    add_heading2("B. Danh sách các Thông điệp (Sequence Messages):")
    table_seq3 = doc.add_table(rows=1, cols=5)
    table_seq3.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = table_seq3.rows[0].cells
    h[0].text = "STT"
    h[1].text = "Gửi từ"
    h[2].text = "Đến"
    h[3].text = "Nội dung Thông điệp"
    h[4].text = "Loại mũi tên"
    for cell in h:
        cell.paragraphs[0].runs[0].font.bold = True

    seq3_data = [
        ("1", "User", "InvoiceCtrl", "1: Truy cập /invoice/{id}", "Message"),
        ("2", "InvoiceCtrl", "BankSvc", "2: getBankSetting()", "Message"),
        ("3", "BankSvc", "InvoiceCtrl", "3: Trả về STK, Ngân hàng, Chủ TK", "Reply Message"),
        ("4", "InvoiceCtrl", "VietQR", "4: Sinh URL mã QR động (Số tiền + HD_ID)", "Message"),
        ("5", "VietQR", "InvoiceCtrl", "5: Trả về ảnh mã VietQR", "Reply Message"),
        ("6", "InvoiceCtrl", "ThermalPage", "6: Render Model (invoice, orderDetails, bankSetting)", "Reply Message"),
        ("7", "ThermalPage", "User", "7: Hiển thị phiếu in nhiệt 80mm kèm Mã QR", "Reply Message"),
        ("8", "User", "ThermalPage", "8: Bấm nút 'In Hóa Đơn (80mm)' (window.print)", "Message")
    ]
    for stt, fr, to, msg, arrow in seq3_data:
        row = table_seq3.add_row().cells
        row[0].text = stt
        row[1].text = fr
        row[2].text = to
        row[3].text = msg
        row[4].text = arrow

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ----------------------------------------------------
    # SƠ ĐỒ 7: ACTIVITY DIAGRAM - POS ORDER FLOW
    # ----------------------------------------------------
    add_heading1("7. SƠ ĐỒ HOẠT ĐỘNG - LUỒNG BÁN HÀNG POS (Activity Diagram - POS Order Flow)")
    
    p = doc.add_paragraph()
    p.add_run("• Mục đích: ").bold = True
    p.add_run("Mô tả quy trình xử lý công việc (Workflow) khi thu ngân thực hiện bán hàng tại quầy.")

    add_heading2("A. Tạo Diagram trong StarUML:")
    doc.add_paragraph("Vào Model -> Add Diagram -> Select 'Activity Diagram'. Naming: 'Act_POS_Order_Flow'.")

    add_heading2("B. Các Node (Hình vẽ) cần tạo:")
    doc.add_paragraph("1. Initial Node (Chấm tròn đen) -> Bắt đầu.")
    doc.add_paragraph("2. Action Node: 'Thu ngân mở màn hình /orders/add'")
    doc.add_paragraph("3. Action Node: 'Chọn món uống từ danh mục / Tìm kiếm'")
    doc.add_paragraph("4. Action Node: 'Tăng/Giảm số lượng & Chọn Voucher'")
    doc.add_paragraph("5. Action Node: 'Chọn Khách hàng lẻ / Thành viên & Số bàn'")
    doc.add_paragraph("6. Action Node: 'Bấm Thanh Toán POS'")
    doc.add_paragraph("7. Decision Node (Hình thoi): 'Kiểm tra món > 0?'")
    doc.add_paragraph("8. Action Node: 'Hiển thị báo lỗi yêu cầu chọn món'")
    doc.add_paragraph("9. Action Node: 'Hệ thống tự động trừ kho nguyên liệu'")
    doc.add_paragraph("10. Action Node: 'Tự động tính & cộng điểm thưởng Khách hàng'")
    doc.add_paragraph("11. Action Node: 'Tạo Hóa đơn & Giao dịch Payment'")
    doc.add_paragraph("12. Action Node: 'Chuyển sang trang In Hóa Đơn Nhiệt 80mm'")
    doc.add_paragraph("13. Activity Final Node (Chấm tròn đen có vòng tròn ngoài) -> Kết thúc.")

    add_heading2("C. Hướng dẫn nối luồng mũi tên (Control Flow):")
    doc.add_paragraph("• Initial Node ➔ Mở màn hình POS ➔ Chọn món ➔ Tăng/Giảm số lượng ➔ Chọn Khách & Bàn ➔ Bấm Thanh Toán ➔ Hình thoi (Kiểm tra món > 0).")
    doc.add_paragraph("• Từ Hình thoi ➔ [Nếu 0 món] ➔ Hiển thị báo lỗi ➔ Quay lại Chọn món.")
    doc.add_paragraph("• Từ Hình thoi ➔ [Nếu > 0 món] ➔ Tự động trừ kho ➔ Tự động cộng điểm ➔ Tạo Hóa đơn ➔ Chuyển sang trang In 80mm ➔ Activity Final Node.")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ----------------------------------------------------
    # SƠ ĐỒ 8: ACTIVITY DIAGRAM - AUTO RECIPE & POINTS
    # ----------------------------------------------------
    add_heading1("8. SƠ ĐỒ HOẠT ĐỘNG - LUỒNG TỰ ĐỘNG TRỪ KHO NGUYÊN LIỆU & TÍCH ĐIỂM")
    
    add_heading2("A. Các Node chính:")
    doc.add_paragraph("1. Initial Node (Bắt đầu xử lý tự động sau checkout).")
    doc.add_paragraph("2. Action Node: 'Lấy danh sách sản phẩm trong đơn'")
    doc.add_paragraph("3. Decision Node: 'Còn sản phẩm trong đơn?'")
    doc.add_paragraph("4. Action Node: 'Tra cứu công thức ProductIngredient'")
    doc.add_paragraph("5. Decision Node: 'Còn nguyên liệu trong món?'")
    doc.add_paragraph("6. Action Node: 'Tính NeededQty = QuantityUsed * Quantity'")
    doc.add_paragraph("7. Action Node: 'Trừ Ingredient.quantity -= NeededQty'")
    doc.add_paragraph("8. Action Node: 'Lưu log InventoryTransaction EXPORT'")
    doc.add_paragraph("9. Decision Node: 'Có chọn Khách hàng thành viên?'")
    doc.add_paragraph("10. Action Node: 'Tính EarnedPoints = TotalAmount / 10000'")
    doc.add_paragraph("11. Action Node: 'Cập nhật Customer.point += EarnedPoints'")
    doc.add_paragraph("12. Final Node (Kết thúc xử lý).")

    add_heading2("B. Hướng dẫn nối luồng lặp (Loop Control Flow):")
    doc.add_paragraph("• Bắt đầu ➔ Lấy danh sách sản phẩm ➔ Hình thoi (Còn sản phẩm?).")
    doc.add_paragraph("• [Còn SP] ➔ Tra cứu công thức ➔ Hình thoi (Còn nguyên liệu?).")
    doc.add_paragraph("• [Còn NL] ➔ Tính NeededQty ➔ Trừ kho ➔ Lưu log EXPORT ➔ Vòng lặp quay lại Hình thoi (Còn nguyên liệu?).")
    doc.add_paragraph("• [Hết NL] ➔ Quay lại Hình thoi (Còn sản phẩm?).")
    doc.add_paragraph("• [Hết SP] ➔ Hình thoi (Có chọn Khách thành viên?).")
    doc.add_paragraph("• [Có Khách] ➔ Tính điểm EarnedPoints ➔ Cộng Customer.point ➔ Kết thúc.")
    doc.add_paragraph("• [Không Khách] ➔ Kết thúc.")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ----------------------------------------------------
    # SƠ ĐỒ 9: STATECHART DIAGRAM - ORDER LIFECYCLE
    # ----------------------------------------------------
    add_heading1("9. SƠ ĐỒ TRẠNG THÁI VÒNG ĐỜI ĐƠN HÀNG (Statechart Diagram - Order Lifecycle)")
    
    p = doc.add_paragraph()
    p.add_run("• Mục đích: ").bold = True
    p.add_run("Mô tả sự chuyển đổi trạng thái của Đơn hàng từ lúc tạo đến khi hoàn tất.")

    add_heading2("A. Tạo Diagram trong StarUML:")
    doc.add_paragraph("Vào Model -> Add Diagram -> Select 'Statechart Diagram'. Naming: 'State_Order_Lifecycle'.")

    add_heading2("B. Các Trạng thái (State Boxes):")
    doc.add_paragraph("1. CART (Giỏ hàng online)")
    doc.add_paragraph("2. PENDING (Đơn chờ duyệt)")
    doc.add_paragraph("3. CREATED (Khởi tạo đơn quầy)")
    doc.add_paragraph("4. COMPLETED (Hoàn tất đơn hàng)")
    doc.add_paragraph("5. CANCELLED (Đã hủy đơn)")
    doc.add_paragraph("6. PAID (Đã phát hành hóa đơn & thanh toán)")

    add_heading2("C. Hướng dẫn nối đường chuyển trạng thái (Transitions):")
    table_st1 = doc.add_table(rows=1, cols=4)
    table_st1.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = table_st1.rows[0].cells
    h[0].text = "Trạng thái đầu (From)"
    h[1].text = "Trạng thái sau (To)"
    h[2].text = "Sự kiện kích hoạt (Trigger Event)"
    h[3].text = "Hành động (Action)"
    for cell in h:
        cell.paragraphs[0].runs[0].font.bold = True

    st1_data = [
        ("Initial State", "CART", "Khách thêm món vào giỏ", "Tạo mới CustomerOrder"),
        ("CART", "PENDING", "Khách bấm 'Đặt hàng Online'", "Cập nhật status PENDING"),
        ("PENDING", "COMPLETED", "Admin/Staff bấm 'Approve'", "Tự động trừ kho & cộng điểm"),
        ("PENDING", "CANCELLED", "Admin/Staff bấm 'Cancel'", "Hủy đơn hàng"),
        ("Initial State", "CREATED", "Thu ngân chọn món tại quầy POS", "Tạo mới Orders"),
        ("CREATED", "COMPLETED", "Thu ngân bấm 'Thanh Toán POS'", "Trừ kho & cộng điểm"),
        ("COMPLETED", "PAID", "Tự động xuất Hóa đơn & Payment", "Hoàn tất vòng đời đơn"),
        ("PAID / CANCELLED", "Final State", "Kết thúc vòng đời", "Lưu trữ lịch sử")
    ]
    for fr, to, tr, ac in st1_data:
        row = table_st1.add_row().cells
        row[0].text = fr
        row[1].text = to
        row[2].text = tr
        row[3].text = ac

    add_callout("Dùng công cụ 'Transition' kéo từ State này sang State kia và điền chữ ở mục Trigger!")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ----------------------------------------------------
    # SƠ ĐỒ 10: STATECHART DIAGRAM - TABLE LIFECYCLE
    # ----------------------------------------------------
    add_heading1("10. SƠ ĐỒ TRẠNG THÁI QUẢN LÝ BÀN ĂN (Statechart Diagram - Table Lifecycle)")
    
    add_heading2("A. Các Trạng thái (State Boxes):")
    doc.add_paragraph("1. EMPTY (Bàn trống / Sẵn sàng)")
    doc.add_paragraph("2. OCCUPIED (Đang có khách ngồi)")
    doc.add_paragraph("3. OUT_OF_SERVICE (Tạm ngưng phục vụ / Bảo trì)")

    add_heading2("B. Hướng dẫn nối đường chuyển trạng thái (Transitions):")
    table_st2 = doc.add_table(rows=1, cols=4)
    table_st2.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = table_st2.rows[0].cells
    h[0].text = "Trạng thái đầu (From)"
    h[1].text = "Trạng thái sau (To)"
    h[2].text = "Sự kiện kích hoạt (Trigger Event)"
    h[3].text = "Hành động (Action)"
    for cell in h:
        cell.paragraphs[0].runs[0].font.bold = True

    st2_data = [
        ("Initial State", "EMPTY", "Khởi tạo bàn mới trong hệ thống", "Tạo mới TableCafe"),
        ("EMPTY", "OCCUPIED", "Gán số bàn tại quầy POS / Khách quét QR", "Chuyển status OCCUPIED"),
        ("OCCUPIED", "EMPTY", "Thu ngân bấm 'Thanh Toán' (/table-payment/{id})", "Giải phóng bàn về EMPTY"),
        ("EMPTY", "OUT_OF_SERVICE", "Thu ngân / Admin bấm ngưng phục vụ", "Bàn bảo trì"),
        ("OUT_OF_SERVICE", "EMPTY", "Thu ngân / Admin mở lại bàn", "Bàn hoạt động lại")
    ]
    for fr, to, tr, ac in st2_data:
        row = table_st2.add_row().cells
        row[0].text = fr
        row[1].text = to
        row[2].text = tr
        row[3].text = ac

    # Footer note
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p_ft = doc.add_paragraph()
    p_ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ft = p_ft.add_run("--- HẾT CẨM NANG HƯỚNG DẪN VẼ STARUML - CHÚC BẠN VẼ SƠ ĐỒ THÀNH CÔNG! ---")
    r_ft.font.bold = True
    r_ft.font.size = Pt(10)
    r_ft.font.color.rgb = RGBColor(139, 94, 60)

    # Save to file
    file_path = "Huong_Dan_Ve_10_So_Do_StarUML_Co_Dao_Quan_POS.docx"
    doc.save(file_path)
    print(f"File created successfully at: {file_path}")

if __name__ == "__main__":
    create_word_guide()
